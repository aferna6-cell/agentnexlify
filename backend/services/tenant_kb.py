"""Tenant KB sources — extract, dedupe, and compile per-tenant knowledge.

One ingest spine serves three transports (specs/drive-kb-onboarding_spec.md):
  - dashboard bulk upload            (source='upload')
  - Google Drive daily sync          (source='drive')
  - the local folder-sync CLI        (source='local_sync')

Flow: raw file bytes -> extract_text (pdf/docx/txt/md) -> pii_scan ->
upsert_document (sha256 diff, unchanged docs skipped) -> compile_tenant_kb
(assemble every active document into widget_configs.knowledge_base with
``<!-- source: ... | synced: ... -->`` provenance headers). The widget and
voice prompts already ground on widget_configs.knowledge_base (migration 077),
so a compile makes new knowledge live without any other change.

client_id (NOT tenant_id) on tenant_kb_documents — customer-data family.
Do NOT add 'from __future__ import annotations' — breaks Pydantic on FastAPI.
"""

import hashlib
import io
import logging
import re
from datetime import datetime, timezone

from backend.models.database import get_service_supabase

logger = logging.getLogger(__name__)

# Doc-count allowance per plan tier (spec Q2/risks: Free 10 / mid 100 / top uncapped)
PLAN_DOC_LIMITS = {
    "free": 10,
    "chatbot": 100,
    "growth": 100,
    "autopilot": 100,
}
DEFAULT_DOC_LIMIT = 1000  # agent_os / professional / enterprise

# Compile cap keeps the merged KB prompt-injectable; the widget prompt
# builder applies its own tighter truncation downstream.
COMPILE_CHAR_CAP = 120_000
MAX_FILE_BYTES = 5 * 1024 * 1024  # spec: skip docs >5MB

SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".txt", ".md", ".markdown")


class UnsupportedDocument(Exception):
    """Raised when a file type can't be converted to text."""


def doc_limit_for_plan(plan: str | None) -> int:
    return PLAN_DOC_LIMITS.get((plan or "").strip().lower(), DEFAULT_DOC_LIMIT)


# ---------------------------------------------------------------------------
# Extraction
# ---------------------------------------------------------------------------


def extract_text(filename: str, data: bytes) -> str:
    """Convert a document to plain markdown-ish text.

    PDF via pypdf, DOCX via python-docx, TXT/MD passthrough. Raises
    UnsupportedDocument for anything else (spec: .pages + scanned PDFs are
    skipped with a logged reason; a scanned PDF surfaces here as empty text).
    """
    if len(data) > MAX_FILE_BYTES:
        raise UnsupportedDocument("file exceeds the 5MB limit")

    name = (filename or "").lower()
    if name.endswith((".txt", ".md", ".markdown")):
        return data.decode("utf-8", errors="replace").strip()

    if name.endswith(".pdf"):
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(data))
            pages = [(page.extract_text() or "") for page in reader.pages]
            text = "\n\n".join(p.strip() for p in pages if p.strip()).strip()
        except UnsupportedDocument:
            raise
        except Exception as exc:
            raise UnsupportedDocument(f"PDF parse failed: {type(exc).__name__}") from exc
        if not text:
            raise UnsupportedDocument("PDF contains no extractable text (scanned? OCR is v2)")
        return text

    if name.endswith(".docx"):
        try:
            import docx

            document = docx.Document(io.BytesIO(data))
            text = "\n\n".join(
                p.text.strip() for p in document.paragraphs if p.text.strip()
            ).strip()
        except Exception as exc:
            raise UnsupportedDocument(f"DOCX parse failed: {type(exc).__name__}") from exc
        if not text:
            raise UnsupportedDocument("DOCX contains no extractable text")
        return text

    raise UnsupportedDocument(f"unsupported file type: {filename}")


# ---------------------------------------------------------------------------
# PII scan (spec Q9 D: regex pre-filter, flag counts, never block)
# ---------------------------------------------------------------------------

_SSN_RE = re.compile(r"\b\d{3}-\d{2}-\d{4}\b")
_CC_RE = re.compile(r"\b(?:\d[ -]?){13,16}\b")


def pii_scan(text: str) -> int:
    """Count of potential PII hits (SSN / credit-card shaped). Advisory only."""
    return len(_SSN_RE.findall(text)) + len(_CC_RE.findall(text))


# ---------------------------------------------------------------------------
# Upsert + compile
# ---------------------------------------------------------------------------


def _sha256(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def count_active_documents(client_id: str) -> int:
    db = get_service_supabase()
    result = (
        db.table("tenant_kb_documents")
        .select("id", count="exact")
        .eq("client_id", client_id)
        .eq("status", "active")
        .execute()
    )
    return result.count or 0


def upsert_document(
    client_id: str,
    *,
    source: str,
    external_id: str,
    filename: str,
    content_md: str,
) -> str:
    """Insert or update one KB document. Returns 'added'|'updated'|'unchanged'.

    Unchanged is decided by content_sha256 — the spec's diff-based recompile
    contract, applied per document.
    """
    db = get_service_supabase()
    digest = _sha256(content_md)
    now_iso = datetime.now(timezone.utc).isoformat()

    existing = (
        db.table("tenant_kb_documents")
        .select("id, content_sha256, status")
        .eq("client_id", client_id)
        .eq("source", source)
        .eq("external_id", external_id)
        .limit(1)
        .execute()
    )
    row = existing.data[0] if existing.data else None

    if row and row.get("content_sha256") == digest and row.get("status") == "active":
        return "unchanged"

    payload = {
        "filename": filename,
        "content_md": content_md,
        "content_sha256": digest,
        "pii_flags": pii_scan(content_md),
        "status": "active",
        "error": None,
        "synced_at": now_iso,
        "updated_at": now_iso,
    }
    if row:
        db.table("tenant_kb_documents").update(payload).eq("id", row["id"]).execute()
        return "updated"

    db.table("tenant_kb_documents").insert(
        {
            "client_id": client_id,
            "source": source,
            "external_id": external_id,
            **payload,
        }
    ).execute()
    return "added"


def compile_tenant_kb(client_id: str) -> dict:
    """Assemble every active document into widget_configs.knowledge_base.

    Each document contributes a provenance header (spec Q5 C):
        <!-- source: drive/pricing.pdf | synced: 2026-07-13 -->
    so the second-brain view and any human reading the KB can trace every
    answer back to its file. Returns {documents, chars, truncated}.
    """
    db = get_service_supabase()
    result = (
        db.table("tenant_kb_documents")
        .select("source, external_id, filename, content_md, synced_at")
        .eq("client_id", client_id)
        .eq("status", "active")
        .order("filename", desc=False)
        .execute()
    )
    docs = result.data or []

    sections: list[str] = []
    total = 0
    truncated = False
    for doc in docs:
        synced_day = (doc.get("synced_at") or "")[:10]
        header = (
            f"<!-- source: {doc.get('source')}/{doc.get('filename')} "
            f"| synced: {synced_day} -->"
        )
        body = (doc.get("content_md") or "").strip()
        section = f"{header}\n\n{body}"
        if total + len(section) > COMPILE_CHAR_CAP:
            truncated = True
            break
        sections.append(section)
        total += len(section)

    merged = "\n\n---\n\n".join(sections)

    db.table("widget_configs").update({"knowledge_base": merged}).eq(
        "tenant_id", client_id
    ).execute()

    logger.info(
        "tenant_kb: compiled %d documents (%d chars%s) for tenant %s",
        len(sections),
        total,
        ", truncated" if truncated else "",
        client_id,
    )

    # Golden-question regression check on the freshly compiled corpus
    # (enterprise-audit item 2). Best-effort — never breaks a compile.
    from backend.services.kb_evals import run_evals_after_compile

    run_evals_after_compile(client_id)

    # Retrievable projection. Fail-open — never breaks compile.
    from backend.services.tenant_kb_index import index_after_compile

    index_after_compile(client_id)

    return {"documents": len(sections), "chars": total, "truncated": truncated}


def ingest_file(
    client_id: str,
    *,
    source: str,
    external_id: str,
    filename: str,
    data: bytes,
) -> dict:
    """Extract + upsert one raw file. Returns a per-file result dict.

    Never raises for per-file problems — callers batch many files and one
    bad PDF must not fail the rest. The caller runs compile_tenant_kb once
    after the batch.
    """
    try:
        text = extract_text(filename, data)
    except UnsupportedDocument as exc:
        return {"filename": filename, "status": "skipped", "reason": str(exc)}
    except Exception as exc:  # defensive: parser libs can throw anything
        logger.warning(
            "tenant_kb: extraction crashed for %s (tenant %s)",
            filename,
            client_id,
            exc_info=True,
        )
        return {"filename": filename, "status": "skipped", "reason": f"parse error: {type(exc).__name__}"}

    try:
        outcome = upsert_document(
            client_id,
            source=source,
            external_id=external_id,
            filename=filename,
            content_md=text,
        )
    except Exception:
        logger.exception(
            "tenant_kb: upsert failed for %s (tenant %s)", filename, client_id
        )
        return {"filename": filename, "status": "error", "reason": "database write failed"}

    return {"filename": filename, "status": outcome, "chars": len(text)}
