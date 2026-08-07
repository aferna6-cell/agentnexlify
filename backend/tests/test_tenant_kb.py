"""Tests for the tenant KB spine (backend/services/tenant_kb.py) and the
documents router (backend/routers/tenant_kb.py) — goal items 1 (bulk upload)
and the shared ingest used by items 2 (Drive) and 4 (local sync CLI).

Contracts:
  - extraction handles md/txt passthrough, real DOCX round-trip, and rejects
    unsupported/oversized/unparseable files with a reason
  - PII scan counts SSN/CC shapes, never blocks
  - upsert is sha-diffed: added -> unchanged -> updated
  - compile writes provenance-headed merged KB into widget_configs (client_id
    tables per schema discipline) and enforces the char cap
  - upload endpoint batches per-file results, respects plan doc limits,
    requires auth + tenant match
"""

import io
from unittest.mock import MagicMock, patch

from backend.services import tenant_kb as tk
from backend.tests.conftest import _make_auth_token

TENANT_ID = "00000000-0000-0000-0000-000000000001"
BASE = f"/api/v1/kb/{TENANT_ID}"


def _auth_headers(tenant_id=TENANT_ID):
    return {"Authorization": f"Bearer {_make_auth_token(tenant_id)}"}


# ---------------------------------------------------------------------------
# Extraction + PII
# ---------------------------------------------------------------------------


class TestExtractText:
    def test_markdown_passthrough(self):
        assert tk.extract_text("faq.md", b"# Hours\nOpen 9-5") == "# Hours\nOpen 9-5"

    def test_txt_passthrough(self):
        assert tk.extract_text("notes.TXT", b"plain notes") == "plain notes"

    def test_docx_round_trip(self):
        import docx

        buf = io.BytesIO()
        d = docx.Document()
        d.add_paragraph("Pricing sheet 2026")
        d.add_paragraph("Balayage from $180")
        d.save(buf)
        text = tk.extract_text("pricing.docx", buf.getvalue())
        assert "Pricing sheet 2026" in text
        assert "Balayage from $180" in text

    def test_unsupported_extension_rejected(self):
        try:
            tk.extract_text("photo.png", b"\x89PNG")
            assert False, "should have raised"
        except tk.UnsupportedDocument as exc:
            assert "unsupported" in str(exc)

    def test_oversized_rejected(self):
        try:
            tk.extract_text("big.md", b"x" * (tk.MAX_FILE_BYTES + 1))
            assert False, "should have raised"
        except tk.UnsupportedDocument as exc:
            assert "5MB" in str(exc)

    def test_malformed_pdf_rejected_with_reason(self):
        try:
            tk.extract_text("bad.pdf", b"not a pdf at all")
            assert False, "should have raised"
        except tk.UnsupportedDocument as exc:
            assert "PDF" in str(exc)


class TestPiiScan:
    def test_counts_ssn_and_card_shapes(self):
        text = "SSN 123-45-6789, card 4111 1111 1111 1111, phone (914) 555-0000"
        assert tk.pii_scan(text) == 2

    def test_clean_text_zero(self):
        assert tk.pii_scan("We open at 9 and close at 5.") == 0


class TestDocLimits:
    def test_tiers(self):
        assert tk.doc_limit_for_plan("free") == 10
        assert tk.doc_limit_for_plan("chatbot") == 100
        assert tk.doc_limit_for_plan("agent_os") == tk.DEFAULT_DOC_LIMIT
        assert tk.doc_limit_for_plan(None) == tk.DEFAULT_DOC_LIMIT


# ---------------------------------------------------------------------------
# Upsert + compile (mocked supabase)
# ---------------------------------------------------------------------------


def _chain(rows):
    result = MagicMock()
    result.data = rows
    result.count = len(rows)
    c = MagicMock()
    for m in ("select", "eq", "neq", "order", "limit", "insert", "update"):
        getattr(c, m).return_value = c
    c.execute.return_value = result
    return c


class TestUpsertDocument:
    def _db(self, existing_rows):
        db = MagicMock()
        db.table.return_value = _chain(existing_rows)
        return db

    def test_new_document_added(self):
        db = self._db([])
        with patch.object(tk, "get_service_supabase", return_value=db):
            outcome = tk.upsert_document(
                TENANT_ID, source="upload", external_id="a.md",
                filename="a.md", content_md="hello",
            )
        assert outcome == "added"

    def test_same_sha_unchanged(self):
        digest = tk._sha256("hello")
        db = self._db([{"id": "d1", "content_sha256": digest, "status": "active"}])
        with patch.object(tk, "get_service_supabase", return_value=db):
            outcome = tk.upsert_document(
                TENANT_ID, source="upload", external_id="a.md",
                filename="a.md", content_md="hello",
            )
        assert outcome == "unchanged"

    def test_changed_content_updated(self):
        db = self._db([{"id": "d1", "content_sha256": "old", "status": "active"}])
        with patch.object(tk, "get_service_supabase", return_value=db):
            outcome = tk.upsert_document(
                TENANT_ID, source="upload", external_id="a.md",
                filename="a.md", content_md="hello v2",
            )
        assert outcome == "updated"


class TestCompile:
    def test_merged_kb_carries_provenance_headers(self):
        docs = [
            {"source": "drive", "external_id": "f1", "filename": "pricing.pdf",
             "content_md": "Balayage $180", "synced_at": "2026-07-13T10:00:00+00:00"},
            {"source": "upload", "external_id": "faq.md", "filename": "faq.md",
             "content_md": "Open 9-5", "synced_at": "2026-07-12T10:00:00+00:00"},
        ]
        db = MagicMock()
        db.table.return_value = _chain(docs)
        with patch.object(tk, "get_service_supabase", return_value=db):
            stats = tk.compile_tenant_kb(TENANT_ID)
        assert stats["documents"] == 2
        written = db.table.return_value.update.call_args.args[0]["knowledge_base"]
        assert "<!-- source: drive/pricing.pdf | synced: 2026-07-13 -->" in written
        assert "<!-- source: upload/faq.md | synced: 2026-07-12 -->" in written
        assert "Balayage $180" in written and "Open 9-5" in written

    def test_char_cap_truncates(self):
        docs = [
            {"source": "upload", "external_id": f"f{i}", "filename": f"f{i}.md",
             "content_md": "x" * 50_000, "synced_at": "2026-07-13T00:00:00+00:00"}
            for i in range(4)
        ]
        db = MagicMock()
        db.table.return_value = _chain(docs)
        with patch.object(tk, "get_service_supabase", return_value=db):
            stats = tk.compile_tenant_kb(TENANT_ID)
        assert stats["truncated"] is True
        assert stats["chars"] <= tk.COMPILE_CHAR_CAP


class TestIngestFile:
    def test_unsupported_file_skipped_not_raised(self):
        outcome = tk.ingest_file(
            TENANT_ID, source="upload", external_id="x.png",
            filename="x.png", data=b"\x89PNG",
        )
        assert outcome["status"] == "skipped"
        assert "unsupported" in outcome["reason"]

    def test_db_failure_reported_not_raised(self):
        with patch.object(tk, "upsert_document", side_effect=RuntimeError("db down")):
            outcome = tk.ingest_file(
                TENANT_ID, source="upload", external_id="a.md",
                filename="a.md", data=b"hello",
            )
        assert outcome["status"] == "error"


# ---------------------------------------------------------------------------
# Router (ASGI + conftest client/mock_supabase)
# ---------------------------------------------------------------------------


class TestUploadEndpoint:
    def test_requires_auth(self, client):
        resp = client.post(f"{BASE}/documents", files={"files": ("a.md", b"hi")})
        assert resp.status_code in (401, 422)

    def test_cross_tenant_forbidden(self, client, mock_supabase):
        resp = client.post(
            f"{BASE}/documents",
            headers=_auth_headers("00000000-0000-0000-0000-000000000002"),
            files={"files": ("a.md", b"hi")},
        )
        assert resp.status_code == 403

    def test_bulk_upload_ingests_and_compiles(self, client, mock_supabase):
        mock_supabase.table.return_value = _chain([])
        files = [
            ("files", ("faq.md", b"# FAQ\nOpen 9-5", "text/markdown")),
            ("files", ("photo.png", b"\x89PNG", "image/png")),
        ]
        resp = client.post(f"{BASE}/documents", headers=_auth_headers(), files=files)
        assert resp.status_code == 200, resp.text
        body = resp.json()
        by_name = {r["filename"]: r for r in body["results"]}
        assert by_name["faq.md"]["status"] == "added"
        assert by_name["photo.png"]["status"] == "skipped"
        assert "compiled" in body

    def test_local_sync_source_recorded(self, client, mock_supabase):
        # The folder-sync CLI posts source=local_sync so provenance is honest
        mock_supabase.table.return_value = _chain([])
        with patch(
            "backend.routers.tenant_kb.ingest_file",
            return_value={"filename": "a.md", "status": "added"},
        ) as ingest:
            resp = client.post(
                f"{BASE}/documents",
                headers=_auth_headers(),
                files={"files": ("a.md", b"hi")},
                data={"source": "local_sync"},
            )
        assert resp.status_code == 200, resp.text
        assert ingest.call_args.kwargs["source"] == "local_sync"

    def test_invalid_source_rejected(self, client, mock_supabase):
        mock_supabase.table.return_value = _chain([])
        resp = client.post(
            f"{BASE}/documents",
            headers=_auth_headers(),
            files={"files": ("a.md", b"hi")},
            data={"source": "dropbox"},
        )
        assert resp.status_code == 400

    def test_plan_doc_limit_enforced(self, client, mock_supabase):
        # Free-plan token; tenant already at the 10-doc limit
        from jose import jwt as jose_jwt
        from datetime import datetime, timedelta, timezone
        from backend.config import settings

        token = jose_jwt.encode(
            {
                "tenant_id": TENANT_ID,
                "email": "o@t.co",
                "plan": "free",
                "business_name": "T",
                "exp": datetime.now(timezone.utc) + timedelta(hours=1),
            },
            settings.api_secret_key,
            algorithm="HS256",
        )
        chain = _chain([])
        chain.execute.return_value.count = 10  # count_active_documents -> 10
        mock_supabase.table.return_value = chain
        resp = client.post(
            f"{BASE}/documents",
            headers={"Authorization": f"Bearer {token}"},
            files={"files": ("more.md", b"more")},
        )
        assert resp.status_code == 200
        assert resp.json()["results"][0]["status"] == "skipped"
        assert "limit" in resp.json()["results"][0]["reason"]


class TestNotesEndpoint:
    def test_requires_auth(self, client):
        resp = client.post(f"{BASE}/notes", json={"title": "Hours", "content": "9-5"})
        assert resp.status_code in (401, 422)

    def test_cross_tenant_forbidden(self, client, mock_supabase):
        resp = client.post(
            f"{BASE}/notes",
            headers=_auth_headers("00000000-0000-0000-0000-000000000002"),
            json={"title": "Hours", "content": "9-5"},
        )
        assert resp.status_code == 403

    def test_empty_content_rejected(self, client, mock_supabase):
        resp = client.post(
            f"{BASE}/notes",
            headers=_auth_headers(),
            json={"title": "Hours", "content": ""},
        )
        assert resp.status_code == 422  # Pydantic min_length

    def test_whitespace_only_rejected(self, client, mock_supabase):
        resp = client.post(
            f"{BASE}/notes",
            headers=_auth_headers(),
            json={"title": "   ", "content": "  \n "},
        )
        assert resp.status_code == 400

    def test_note_saved_as_note_source_and_compiles(self, client, mock_supabase):
        mock_supabase.table.return_value = _chain([])
        with patch(
            "backend.routers.tenant_kb.upsert_document", return_value="added"
        ) as upsert:
            resp = client.post(
                f"{BASE}/notes",
                headers=_auth_headers(),
                json={"title": "Cancellation Policy!", "content": "24h notice."},
            )
        assert resp.status_code == 200, resp.text
        body = resp.json()
        assert body["status"] == "added"
        assert body["filename"] == "cancellation-policy.md"
        assert "compiled" in body
        kwargs = upsert.call_args.kwargs
        assert kwargs["source"] == "note"
        assert kwargs["external_id"] == "cancellation-policy"
        # Title becomes a markdown heading so the compiled KB stays readable
        assert kwargs["content_md"] == "# Cancellation Policy!\n\n24h notice."

    def test_new_note_blocked_at_plan_limit(self, client, mock_supabase):
        # Free-plan token; tenant already at the 10-doc limit; no existing note
        from jose import jwt as jose_jwt
        from datetime import datetime, timedelta, timezone
        from backend.config import settings

        token = jose_jwt.encode(
            {
                "tenant_id": TENANT_ID,
                "email": "o@t.co",
                "plan": "free",
                "business_name": "T",
                "exp": datetime.now(timezone.utc) + timedelta(hours=1),
            },
            settings.api_secret_key,
            algorithm="HS256",
        )
        chain = _chain([])
        chain.execute.return_value.count = 10  # count_active_documents -> 10
        mock_supabase.table.return_value = chain
        resp = client.post(
            f"{BASE}/notes",
            headers={"Authorization": f"Bearer {token}"},
            json={"title": "One more", "content": "text"},
        )
        assert resp.status_code == 400
        assert "limit" in resp.json()["detail"]

    def test_existing_note_updates_even_at_limit(self, client, mock_supabase):
        # Same title already stored -> update path skips the limit check
        chain = _chain([{"id": "d1"}])
        chain.execute.return_value.count = 10
        mock_supabase.table.return_value = chain
        with patch(
            "backend.routers.tenant_kb.upsert_document", return_value="updated"
        ):
            resp = client.post(
                f"{BASE}/notes",
                headers=_auth_headers(),
                json={"title": "Cancellation Policy!", "content": "48h notice now."},
            )
        assert resp.status_code == 200, resp.text
        assert resp.json()["status"] == "updated"


class TestNoteSlug:
    def test_slug_shapes(self):
        from backend.routers.tenant_kb import _note_external_id

        assert _note_external_id("Cancellation Policy!") == "cancellation-policy"
        assert _note_external_id("  Hours & Parking  ") == "hours-parking"
        assert _note_external_id("???") == "note"
        assert len(_note_external_id("x" * 500)) <= 120


class TestListAndDelete:
    def test_list_documents_with_provenance(self, client, mock_supabase):
        rows = [
            {"id": "d1", "source": "drive", "filename": "pricing.pdf",
             "status": "active", "pii_flags": 0, "content_md": "abc",
             "synced_at": "2026-07-13T00:00:00+00:00"},
        ]
        mock_supabase.table.return_value = _chain(rows)
        resp = client.get(f"{BASE}/documents", headers=_auth_headers())
        assert resp.status_code == 200
        body = resp.json()
        assert body["total"] == 1
        assert body["documents"][0]["source"] == "drive"
        assert body["documents"][0]["chars"] == 3

    def test_delete_missing_404(self, client, mock_supabase):
        mock_supabase.table.return_value = _chain([])
        resp = client.delete(f"{BASE}/documents/nope", headers=_auth_headers())
        assert resp.status_code == 404

    def test_delete_recompiles(self, client, mock_supabase):
        mock_supabase.table.return_value = _chain([{"id": "d1"}])
        resp = client.delete(f"{BASE}/documents/d1", headers=_auth_headers())
        assert resp.status_code == 200
        assert resp.json()["deleted"] is True


# ---------------------------------------------------------------------------
# Sync token — long-lived scoped auth for the folder-sync CLI
# ---------------------------------------------------------------------------

from datetime import datetime, timedelta, timezone  # noqa: E402

from jose import jwt as jose_jwt  # noqa: E402

from backend.config import settings  # noqa: E402

DRIVE_STATUS = "/api/v1/kb/integrations/drive/status"


def _sync_token(tenant_id=TENANT_ID, purpose="kb_sync", days=180):
    return jose_jwt.encode(
        {
            "tenant_id": tenant_id,
            "purpose": purpose,
            "exp": datetime.now(timezone.utc) + timedelta(days=days),
        },
        settings.api_secret_key,
        algorithm="HS256",
    )


def _tables(mapping):
    """table.side_effect router: table name -> chain (default empty)."""

    def route(name):
        return mapping.get(name) or _chain([])

    return route


class TestSyncToken:
    def test_mint_returns_scoped_long_lived_token(self, client, mock_supabase):
        resp = client.post(f"{BASE}/sync-token", headers=_auth_headers())
        assert resp.status_code == 200, resp.text
        body = resp.json()
        claims = jose_jwt.decode(
            body["token"], settings.api_secret_key, algorithms=["HS256"]
        )
        assert claims["purpose"] == "kb_sync"
        assert claims["tenant_id"] == TENANT_ID
        assert "plan" not in claims  # plan resolved live, never baked in
        remaining = datetime.fromtimestamp(claims["exp"], tz=timezone.utc) - datetime.now(
            timezone.utc
        )
        assert remaining > timedelta(days=170)
        assert body["expires_in_days"] == 180

    def test_mint_requires_matching_tenant(self, client, mock_supabase):
        other = "00000000-0000-0000-0000-000000000002"
        resp = client.post(
            f"/api/v1/kb/{other}/sync-token", headers=_auth_headers()
        )
        assert resp.status_code == 403

    def test_sync_token_lists_documents_with_live_plan(self, client, mock_supabase):
        mock_supabase.table.side_effect = _tables(
            {"tenants": _chain([{"plan": "free"}])}
        )
        try:
            resp = client.get(
                f"{BASE}/documents",
                headers={"Authorization": f"Bearer {_sync_token()}"},
            )
        finally:
            mock_supabase.table.side_effect = None
        assert resp.status_code == 200, resp.text
        assert resp.json()["limit"] == 10  # free-plan limit came from the DB

    def test_sync_token_uploads(self, client, mock_supabase):
        mock_supabase.table.return_value = _chain([])
        resp = client.post(
            f"{BASE}/documents",
            headers={"Authorization": f"Bearer {_sync_token()}"},
            files={"files": ("faq.md", b"# Hours")},
            data={"source": "local_sync"},
        )
        assert resp.status_code == 200, resp.text
        assert resp.json()["results"][0]["status"] == "added"

    def test_sync_token_rejected_outside_kb_router(self, client, mock_supabase):
        # A leaked 180-day token must not open the rest of the dashboard
        resp = client.get(
            DRIVE_STATUS, headers={"Authorization": f"Bearer {_sync_token()}"}
        )
        assert resp.status_code == 401

    def test_sync_token_cannot_mint_another(self, client, mock_supabase):
        resp = client.post(
            f"{BASE}/sync-token",
            headers={"Authorization": f"Bearer {_sync_token()}"},
        )
        assert resp.status_code == 401

    def test_wrong_purpose_token_rejected_on_kb_endpoints(self, client, mock_supabase):
        token = _sync_token(purpose="drive_kb_oauth")
        resp = client.get(
            f"{BASE}/documents", headers={"Authorization": f"Bearer {token}"}
        )
        assert resp.status_code == 401

    def test_cross_tenant_sync_token_forbidden(self, client, mock_supabase):
        token = _sync_token(tenant_id="00000000-0000-0000-0000-000000000002")
        resp = client.get(
            f"{BASE}/documents", headers={"Authorization": f"Bearer {token}"}
        )
        assert resp.status_code == 403
